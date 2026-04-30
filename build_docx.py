"""
Render REPORT.docx from data/runs.json + report_assets/*.png.
Drag the resulting file into Google Drive to open as a native Google Doc.
"""
import json
from collections import OrderedDict, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
RUNS = json.loads((ROOT / "data" / "runs.json").read_text())
ASSETS = ROOT / "report_assets"
OUT = ROOT / "REPORT.docx"

BP_SIZES = [10, 30, 50, 70, 90, 110]
VU_STEPS = [10, 20, 40, 80, 160, 320]

ENGINES = OrderedDict([
    ("maria122", {"display": "MariaDB 12.2.2"}),
    ("maria123", {"display": "MariaDB 12.3.1"}),
    ("mysql84",  {"display": "MySQL 8.4.8"}),
    ("mysql97",  {"display": "MySQL 9.7.0"}),
])
EIDS = list(ENGINES.keys())
MYSQL_EIDS = {"mysql84", "mysql97"}


def _instances_ok(r):
    return r.get("bp_instances") == 2 if r["engine_id"] in MYSQL_EIDS else True


def bp_sweep_run(eid):
    by_run = defaultdict(set)
    for r in RUNS:
        if r["engine_id"] != eid or r["vu"] != 80 or not _instances_ok(r):
            continue
        by_run[r["run_dir"]].add(r["bp_gib"])
    full = [(rd, s) for rd, s in by_run.items() if set(BP_SIZES).issubset(s)]
    return sorted(full, key=lambda x: x[0])[-1][0] if full else None


def bp_record(eid, bp):
    rd = bp_sweep_run(eid)
    if not rd:
        return None
    for r in RUNS:
        if (r["run_dir"] == rd and r["bp_gib"] == bp and r["vu"] == 80
                and _instances_ok(r)):
            return r
    return None


def vu_record(eid, vu):
    cands = [r for r in RUNS
             if r["engine_id"] == eid and r["bp_gib"] == 110
             and r["vu"] == vu and r["tpm"] and _instances_ok(r)]
    return sorted(cands, key=lambda r: r["timestamp_utc"] or "")[-1] if cands else None


def run_tpm(r):
    """HammerDB-reported TPM if present, else the per-second series mean."""
    if not r or not r.get("tpm") or not r["tpm"].get("avg"):
        return None
    hdb = r.get("hammerdb_reported")
    if hdb and hdb.get("tpm"):
        return float(hdb["tpm"])
    return r["tpm"]["avg"]


def scaled_stats(r):
    """Mean/std/p5/p95/CV directly from the native-TPM summary in runs.json."""
    if not r or not r.get("tpm"):
        return None
    t = r["tpm"]
    return {
        "mean": t["avg"],
        "std":  t["std"],
        "p5":   t["p5"],
        "p95":  t["p95"],
        "cv_pct": t["cv_pct"],
    }


def fmt_int(n):
    return f"{int(round(n)):,}"


# ── cell shading helper ──────────────────────────────────────────────────────
def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def style_header(row):
    for cell in row.cells:
        _shade(cell, "1F2937")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xEE, 0xF0, 0xF8)
                run.font.size = Pt(10)


def bold_winner_cells(row, winners_idx):
    """Bold the cell whose column index is in winners_idx."""
    for i, cell in enumerate(row.cells):
        if i in winners_idx:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


# ── sections ─────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return h


def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_image(doc, path, width_in=6.5):
    doc.add_picture(str(path), width=Inches(width_in))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def build_table(doc, headers, rows, winner_cols_per_row=None):
    """
    rows: list of tuples of cells (strings).
    winner_cols_per_row: list[set[int]] same length as rows; indices (relative to
      full row including label col) that should be bolded.
    """
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
    style_header(tbl.rows[0])

    for r_idx, row_vals in enumerate(rows):
        trow = tbl.rows[r_idx + 1]
        for i, v in enumerate(row_vals):
            trow.cells[i].text = str(v)
            for p in trow.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if winner_cols_per_row:
            bold_winner_cells(trow, winner_cols_per_row[r_idx])
    return tbl


# ── main ─────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # default font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── title + banner ──
    title = doc.add_heading("Database Benchmark Comparison — TPROC-C Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    for line in [
        "HammerDB 5.0 · TPROC-C · 1000 warehouses · 3600 s runs · 600 s ramp-up",
        "Hardware: beast-node2.tp.int.percona.com · 80 logical CPUs · 187.54 GiB RAM · NVMe",
        "OS: Ubuntu 24.04 · kernel 6.8.0-60-generic · governor=performance · THP=off · swappiness=1",
        f"Engines: MariaDB 12.2.2, MariaDB 12.3.1-rc, MySQL 8.4.8, MySQL 9.7.0  ·  Generated: {datetime.now():%Y-%m-%d}",
    ]:
        add_para(doc, line, bold=True, size=10)

    # ── Executive Summary ──
    add_heading(doc, "Executive Summary", level=1)
    peak_bp = []
    for eid in EIDS:
        vals = [run_tpm(bp_record(eid, bp)) for bp in BP_SIZES]
        vals = [v for v in vals if v is not None]
        peak_bp.append(max(vals) if vals else None)
    peak_vu = []
    for eid in EIDS:
        vals = [run_tpm(vu_record(eid, vu)) for vu in VU_STEPS]
        vals = [v for v in vals if v is not None]
        peak_vu.append(max(vals) if vals else None)
    scaling = []
    for eid in EIDS:
        lo = run_tpm(vu_record(eid, 10))
        hi = None
        for vu in reversed(VU_STEPS):
            hi = run_tpm(vu_record(eid, vu))
            if hi:
                break
        scaling.append(hi / lo if (lo and hi) else None)

    exec_headers = ["Config"] + [ENGINES[e]["display"] for e in EIDS]
    rows = [
        ("Peak TPM (BP iterations, 80 VU)",
         *[fmt_int(v) if v else "—" for v in peak_bp]),
        ("Peak TPM (VU iterations, BP 110G)",
         *[fmt_int(v) if v else "—" for v in peak_vu]),
        ("Scaling 10→320 VU (BP 110G)",
         *[f"{v:.1f}×" if v else "—" for v in scaling]),
    ]
    build_table(doc, exec_headers, rows)

    best_bp_idx = max(range(len(EIDS)), key=lambda i: peak_bp[i] or 0)
    best_vu_idx = max(range(len(EIDS)), key=lambda i: peak_vu[i] or 0)
    add_para(doc,
             f"Headline: {ENGINES[EIDS[best_bp_idx]]['display']} posts the highest "
             f"BP-iterations peak at 80 VU; {ENGINES[EIDS[best_vu_idx]]['display']} "
             "leads the VU-iterations peak at BP 110 GiB. TPM comes from HammerDB's own "
             "per-second hdbtcount_*.log; steady-state mean confirmed against the TEST RESULT "
             "line in hammerdbcli.out. MySQL columns restricted to runs with "
             "innodb_buffer_pool_instances = 2 for apples-to-apples comparison.",
             size=10)

    # ── Buffer Pool Iterations ──
    add_heading(doc, "Buffer Pool Iterations — 80 VU, 10G–110G", level=1)
    add_para(doc,
             "The InnoDB Buffer Pool caches table and index pages in memory. Iterations vary "
             "innodb_buffer_pool_size from 10 GiB to 110 GiB while holding everything else "
             "constant — 80 virtual users, 1000 warehouses (~100 GB working set), same "
             "hardware, same configuration.")
    add_image(doc, ASSETS / "fig1_bp_line.png")
    add_caption(doc, "Fig 1 — TPROC-C Throughput vs Buffer Pool Size (80 VU · 60-min runs)")

    bp_headers = ["BP Size"] + [ENGINES[e]["display"] for e in EIDS]
    bp_rows, bp_winners = [], []
    for bp in BP_SIZES:
        tpms = [run_tpm(bp_record(eid, bp)) for eid in EIDS]
        vals_idx = {i + 1: v for i, v in enumerate(tpms) if v is not None}
        mx = max(vals_idx.values()) if vals_idx else None
        winners = {i for i, v in vals_idx.items() if v == mx}
        bp_rows.append((f"{bp}G", *[fmt_int(v) if v else "—" for v in tpms]))
        bp_winners.append(winners)
    build_table(doc, bp_headers, bp_rows, bp_winners)

    # ── Virtual Users Iterations ──
    add_heading(doc, "Virtual Users Iterations — BP 110G, 10–320 VU", level=1)
    add_para(doc,
             "A Virtual User (VU) simulates an independent database client. Concurrency was "
             "iterated through {10, 20, 40, 80, 160, 320} virtual users with a fixed 110 GiB "
             "buffer pool. MariaDB 12.2.2 was only benchmarked at 80 VU in this dataset.")
    add_image(doc, ASSETS / "fig3_vu_line.png")
    add_caption(doc, "Fig 3 — TPROC-C Throughput vs Concurrency (BP 110G)")
    add_image(doc, ASSETS / "fig4_scaling.png")
    add_caption(doc, "Fig 4 — Concurrency Scaling Efficiency (BP 110G)")

    vu_headers = ["VU"] + [ENGINES[e]["display"] for e in EIDS]
    vu_rows, vu_winners = [], []
    for vu in VU_STEPS:
        tpms = [run_tpm(vu_record(eid, vu)) for eid in EIDS]
        vals_idx = {i + 1: v for i, v in enumerate(tpms) if v is not None}
        mx = max(vals_idx.values()) if vals_idx else None
        winners = {i for i, v in vals_idx.items() if v == mx}
        vu_rows.append((str(vu), *[fmt_int(v) if v else "—" for v in tpms]))
        vu_winners.append(winners)
    build_table(doc, vu_headers, vu_rows, vu_winners)

    # ── TPM Jitter ──
    add_heading(doc, "TPM Jitter — steady-state windows", level=1)
    add_para(doc,
             "TPM Jitter quantifies the spread of per-second throughput variation during "
             "steady-state (ramp-up excluded). CV% = std ÷ mean × 100; lower is more stable. "
             "P5 and P95 are the 5th and 95th percentile per-second TPM.")
    add_image(doc, ASSETS / "fig5_timeseries.png")
    add_caption(doc, "Fig 5 — TPM Over Time at BP 50G / 80 VU (thin = per-second, thick = 60 s rolling)")

    add_heading(doc, "Buffer Pool Iterations", level=2)
    add_image(doc, ASSETS / "fig6_jitter_bp.png")
    add_caption(doc, "Fig 6 — TPM Jitter — Buffer Pool Iterations (boxes = P25–P75, whiskers = P5–P95)")

    jit_headers = ["Config", "Engine", "Mean TPM", "Std Dev", "CV%", "P5", "P95", "P5–P95"]
    jit_rows = []
    for bp in BP_SIZES:
        for eid in EIDS:
            s = scaled_stats(bp_record(eid, bp))
            if not s:
                continue
            jit_rows.append((f"{bp}G", ENGINES[eid]["display"],
                             fmt_int(s["mean"]), fmt_int(s["std"]),
                             f"{s['cv_pct']:.1f}%",
                             fmt_int(s["p5"]), fmt_int(s["p95"]),
                             fmt_int(s["p95"] - s["p5"])))
    build_table(doc, jit_headers, jit_rows)

    add_heading(doc, "Virtual Users Iterations", level=2)
    add_image(doc, ASSETS / "fig7_jitter_vu.png")
    add_caption(doc, "Fig 7 — TPM Jitter — Virtual Users Iterations (boxes = P25–P75, whiskers = P5–P95)")

    vu_jit_rows = []
    for vu in VU_STEPS:
        for eid in EIDS:
            s = scaled_stats(vu_record(eid, vu))
            if not s:
                continue
            vu_jit_rows.append((f"{vu} VU", ENGINES[eid]["display"],
                                fmt_int(s["mean"]), fmt_int(s["std"]),
                                f"{s['cv_pct']:.1f}%",
                                fmt_int(s["p5"]), fmt_int(s["p95"]),
                                fmt_int(s["p95"] - s["p5"])))
    build_table(doc, jit_headers, vu_jit_rows)

    # ── Methodology ──
    add_heading(doc, "Methodology", level=1)
    for line in [
        "Benchmark: TPROC-C via HammerDB 5.0 (hammerdb_run.tcl)",
        "Workload: 1000 warehouses (~100 GB), 600 s ramp-up, 3600 s measurement, partitioned InnoDB",
        "Driver: timed, tc_refresh_seconds=1",
        "Hardware: beast-node2.tp.int.percona.com — 80 logical CPUs, 187.54 GiB RAM",
        "OS: Ubuntu 24.04, kernel 6.8.0-60-generic; governor=performance, THP=off, swappiness=1, CPU idle POLL+C1 only",
        "Metric: TPM — per-second samples from HammerDB's hdbtcount_*.log; end-of-run mean cross-checked against hammerdbcli.out TEST RESULT line",
        "BP iterations: 80 VU, buffer pool ∈ {10, 30, 50, 70, 90, 110} GiB",
        "VU iterations: 110 GiB buffer pool, VU ∈ {10, 20, 40, 80, 160, 320}",
        "MySQL columns restricted to innodb_buffer_pool_instances = 2; MariaDB columns use default instances = 1",
        "Tail trim: trailing ramp-down samples below 10% of steady-state median are dropped",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(line)
        run.font.size = Pt(10)

    add_para(doc,
             "Data source: github.com/Percona-Lab-results/mysql-tpcc-buffer-pool-sweep",
             size=9, color=RGBColor(0x64, 0x74, 0x8B))

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
